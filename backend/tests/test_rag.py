"""Chunking, retrieval scoping, and ingestion orchestration (DESIGN.md §6)."""

from __future__ import annotations

import io
import uuid
from dataclasses import dataclass
from unittest.mock import AsyncMock, MagicMock

import pytest
from app.agent.ingestion import (
    chunk_text,
    extract_text,
    ingest_source,
    parse_csv,
    parse_docx,
)
from app.agent.rag import RetrievedChunk, build_context_block, retrieve
from app.models.enums import KnowledgeSourceType
from app.models.knowledge import KnowledgeSource


# --- chunking ----------------------------------------------------------------- #
def test_chunk_text_respects_paragraph_boundaries():
    text = "para one word " * 5 + "\n\n" + "para two word " * 5
    chunks = chunk_text(text, chunk_tokens=100, overlap_tokens=0)
    assert len(chunks) == 1  # both paragraphs fit under the target easily
    assert "para one" in chunks[0]
    assert "para two" in chunks[0]


def test_chunk_text_splits_when_over_target():
    para_a = " ".join(f"a{i}" for i in range(60))
    para_b = " ".join(f"b{i}" for i in range(60))
    chunks = chunk_text(f"{para_a}\n\n{para_b}", chunk_tokens=60, overlap_tokens=0)
    assert len(chunks) == 2
    assert "a0" in chunks[0] and "b0" not in chunks[0]
    assert "b0" in chunks[1]


def test_chunk_text_overlap_carries_tail_words():
    para_a = " ".join(f"a{i}" for i in range(60))
    para_b = " ".join(f"b{i}" for i in range(60))
    chunks = chunk_text(f"{para_a}\n\n{para_b}", chunk_tokens=60, overlap_tokens=10)
    # The tail of chunk 1 should reappear at the head of chunk 2.
    tail_of_first = chunks[0].split()[-10:]
    head_of_second = chunks[1].split()[:10]
    assert tail_of_first == head_of_second


def test_chunk_text_splits_oversized_single_paragraph():
    huge_para = " ".join(f"w{i}" for i in range(250))
    chunks = chunk_text(huge_para, chunk_tokens=100, overlap_tokens=0)
    assert len(chunks) >= 2
    assert all(len(c.split()) <= 100 for c in chunks)


def test_chunk_text_empty_input():
    assert chunk_text("   \n\n   ") == []


# --- parsers ------------------------------------------------------------------ #
def test_parse_csv_with_header():
    raw = "name,phone\nAli,+92123\nSara,+92456"
    text = parse_csv(raw)
    assert text == "name: Ali, phone: +92123\nname: Sara, phone: +92456"


def test_extract_text_dispatches_by_type():
    assert extract_text(source_type="markdown", raw_text="# hi") == "# hi"
    assert extract_text(source_type="csv", raw_text="a,b\n1,2") == "a: 1, b: 2"


def test_parse_docx_roundtrip():
    import docx

    doc = docx.Document()
    doc.add_paragraph("Hello from a DOCX paragraph.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    text = parse_docx(buf.getvalue())
    assert "Hello from a DOCX paragraph." in text
    assert "Second paragraph." in text


def test_extract_text_pdf_requires_bytes():
    with pytest.raises(ValueError):
        extract_text(source_type="pdf", raw_text="no bytes given")


# --- retrieval scoping ---------------------------------------------------------- #
@dataclass
class _FakeChunk:
    id: uuid.UUID
    source_id: uuid.UUID
    content: str


class _FakeResult:
    def __init__(self, rows):
        self._rows = rows

    def all(self):
        return self._rows


class _FakeEmbedder:
    def __init__(self, vector):
        self._vector = vector

    async def embed(self, texts, *, model=None):
        return [self._vector for _ in texts]


async def test_retrieve_scopes_by_tenant_and_excludes_tombstoned():
    tenant_id = uuid.uuid4()
    chunk = _FakeChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="hello")
    db = AsyncMock()
    db.execute.return_value = _FakeResult([(chunk, 0.1)])

    results = await retrieve(db, tenant_id, "what are your hours?", embedder=_FakeEmbedder([0.1]))

    assert len(results) == 1
    assert results[0].content == "hello"
    assert results[0].score == pytest.approx(0.9)

    stmt = db.execute.await_args.args[0]
    compiled = stmt.compile()
    sql = str(compiled)
    assert "tenant_id" in sql
    assert "tombstoned" in sql
    assert "<=>" in sql  # pgvector cosine distance operator
    assert "LIMIT" in sql
    # The tenant filter is bound as a parameter (never string-interpolated).
    bound_tenant_ids = [v for v in compiled.params.values() if v == tenant_id]
    assert bound_tenant_ids == [tenant_id]


async def test_retrieve_filters_below_min_score():
    db = AsyncMock()
    chunk = _FakeChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="irrelevant")
    db.execute.return_value = _FakeResult([(chunk, 1.9)])  # score = 1 - 1.9 = -0.9

    results = await retrieve(
        db, uuid.uuid4(), "query", embedder=_FakeEmbedder([0.1]), min_score=0.5
    )
    assert results == []


async def test_retrieve_empty_query_short_circuits():
    db = AsyncMock()
    results = await retrieve(db, uuid.uuid4(), "   ", embedder=_FakeEmbedder([0.1]))
    assert results == []
    db.execute.assert_not_awaited()


def test_build_context_block_numbers_chunks():
    chunks = [
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="first", score=0.9),
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="second", score=0.8),
    ]
    block = build_context_block(chunks)
    assert block == "[1] first\n\n[2] second"


def test_build_context_block_empty():
    assert build_context_block([]) == ""


def test_build_context_block_dedupes_near_identical():
    chunks = [
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="Open 9 to 5", score=0.9),
        # Same content, different whitespace/case (ingestion overlap) → dropped.
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="open 9 to 5", score=0.8),
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content="We ship daily", score=0.7),
    ]
    block = build_context_block(chunks)
    assert block == "[1] Open 9 to 5\n\n[2] We ship daily"


def test_build_context_block_respects_token_budget():
    big = "word " * 500  # ~2500 chars each
    chunks = [
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content=big + "a", score=0.9),
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content=big + "b", score=0.8),
        RetrievedChunk(id=uuid.uuid4(), source_id=uuid.uuid4(), content=big + "c", score=0.7),
    ]
    # 100-token budget = 400 chars; only the top chunk fits (always kept).
    block = build_context_block(chunks, max_tokens=100)
    assert block.count("[") == 1
    assert block.startswith("[1] ")


# --- ingestion orchestration (mocked db) --------------------------------------- #
async def test_ingest_source_tombstones_then_inserts():
    tenant_id = uuid.uuid4()
    source = KnowledgeSource(
        id=uuid.uuid4(),
        tenant_id=tenant_id,
        type=KnowledgeSourceType.manual,
        name="FAQ",
    )
    db = AsyncMock()
    added = []
    db.add = MagicMock(side_effect=lambda obj: added.append(obj))

    chunks = await ingest_source(
        db, source, text="Some short knowledge text.", embedder=_FakeEmbedder([0.1, 0.2])
    )

    # First call is the tombstone UPDATE.
    tombstone_stmt = db.execute.await_args_list[0].args[0]
    compiled = str(tombstone_stmt.compile(compile_kwargs={"literal_binds": True}))
    assert "UPDATE" in compiled.upper()
    assert "tombstoned" in compiled

    assert len(chunks) == 1
    assert added == chunks
    assert chunks[0].tenant_id == tenant_id
    assert chunks[0].source_id == source.id
    assert chunks[0].embedding == [0.1, 0.2]


async def test_ingest_source_empty_text_no_chunks():
    source = KnowledgeSource(
        id=uuid.uuid4(), tenant_id=uuid.uuid4(), type=KnowledgeSourceType.manual, name="Empty"
    )
    db = AsyncMock()
    chunks = await ingest_source(db, source, text="   ", embedder=_FakeEmbedder([0.1]))
    assert chunks == []
